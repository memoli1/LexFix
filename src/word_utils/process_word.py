from docx import Document
from language_tool_python import LanguageTool

class WordProcessor:
    """A class for processing Word documents with error checking and formatting preservation."""

    def __init__(self, language='en-US'):
        self.tool = LanguageTool(language)

    def check_errors_in_word(self, file_path):
        """
        Check for errors in a Word document.

        Args:
            file_path (str): Path to the Word document.

        Returns:
            list: A list of errors with their context and suggestions.
        """
        document = Document(file_path)
        errors = []

        # Process each paragraph for errors
        for paragraph in document.paragraphs:
            text = paragraph.text
            if text.strip():
                matches = self.tool.check(text)
                for match in matches:
                    errors.append({
                        "context": match.context,  # Error context in the text
                        "message": match.message,  # Description of the error
                        "suggestions": match.replacements  # Suggestions for correction
                    })
        return errors

    def correct_text_in_word(self, file_path, output_path, selected_errors):
        """
        Correct selected errors in a Word document while preserving formatting.

        Args:
            file_path (str): Path to the input Word file.
            output_path (str): Path to save the corrected Word file.
            selected_errors (list): List of errors to correct.
        """
        document = Document(file_path)

        for paragraph in document.paragraphs:
            original_text = paragraph.text
            if original_text.strip():
                # Apply only selected corrections
                corrections = {}
                for error in selected_errors:
                    if error['context'] in original_text:
                        start = original_text.find(error['context'])
                        end = start + len(error['context'])
                        corrections[(start, end)] = error['suggestions'][0]

                # Update paragraph with corrections
                self._apply_corrections_to_runs(paragraph, corrections)

        document.save(output_path)

    def _apply_corrections_to_runs(self, paragraph, corrections):
        """
        Apply text corrections to paragraph runs while preserving formatting.

        Args:
            paragraph (docx.Paragraph): The paragraph to process.
            corrections (dict): A mapping of text offsets to corrected replacements.
        """
        runs = paragraph.runs
        corrected_text = paragraph.text

        # Apply corrections to the paragraph text
        for (start, end), replacement in sorted(corrections.items(), reverse=True):
            corrected_text = corrected_text[:start] + replacement + corrected_text[end:]

        # Update runs with corrected text while preserving formatting
        current_offset = 0
        for run in runs:
            run_length = len(run.text)
            if current_offset >= len(corrected_text):
                run.text = ""
            else:
                corrected_run_text = corrected_text[current_offset:current_offset + run_length]
                run.text = corrected_run_text
            current_offset += run_length
